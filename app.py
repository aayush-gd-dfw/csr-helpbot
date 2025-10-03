#!/usr/bin/env python3
import os
import json
import uuid
from typing import List, Dict, Tuple
from flask import Flask, request, jsonify, render_template
import numpy as np
import re
import requests
import pandas as pd  # NEW

# --- Embeddings (free, local) ---
from sentence_transformers import SentenceTransformer

# --- FAISS vector index ---
import faiss

# --- Doc loaders ---
from pypdf import PdfReader
import docx  # python-docx

# ==================== CONFIG ====================
GROQ_API_KEY = "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"
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")

TOP_K = int(os.getenv("TOP_K", "5"))
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "300"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "50"))
EMBED_MODEL_NAME = os.getenv("EMBED_MODEL_NAME", "all-MiniLM-L6-v2")

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md", ".csv"}

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
KNOWLEDGE_DIR = os.path.abspath(os.getenv("KNOWLEDGE_DIR", os.path.join(BASE_DIR, "knowledgebase")))
STORAGE_DIR = os.path.abspath(os.getenv("STORAGE_DIR", os.path.join(BASE_DIR, "storage")))
INDEX_PATH = os.path.join(STORAGE_DIR, "faiss.index")
META_PATH = os.path.join(STORAGE_DIR, "meta.jsonl")

# =================================================

app = Flask(__name__)
embed_model = SentenceTransformer(EMBED_MODEL_NAME, cache_folder=os.path.join(BASE_DIR, "models"))

# --------- Utilities ---------
def clean_text(txt: str) -> str:
    return re.sub(r"\s+", " ", txt).strip()

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    text = clean_text(text)
    chunks, start = [], 0
    step = max(chunk_size - overlap, 1)
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end])
        start += step
    return chunks

def embed_texts(texts: List[str]) -> np.ndarray:
    embs = embed_model.encode(
        texts, convert_to_numpy=True, show_progress_bar=False, normalize_embeddings=True
    )
    return embs.astype("float32")

def load_pdf(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def load_docx(path: str) -> str:
    d = docx.Document(path)
    texts = []
    for p in d.paragraphs:
        texts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join([t for t in texts if t.strip()])

def load_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def load_csv(path: str) -> str:
    """Load CSV transcript into plain text"""
    try:
        df = pd.read_csv(path)
    except Exception as e:
        print(f"⚠️ Failed to read CSV {path}: {e}")
        return ""

    texts = []
    # if likely transcript schema
    cols_lower = [c.lower() for c in df.columns]
    if "speaker" in cols_lower and "transcript" in cols_lower:
        for _, row in df.iterrows():
            spk = str(row.get("speaker", "")).strip()
            txt = str(row.get("transcript", "")).strip()
            if txt:
                texts.append(f"{spk}: {txt}" if spk else txt)
    else:
        for _, row in df.iterrows():
            line = " ".join(str(x) for x in row if pd.notna(x))
            if line.strip():
                texts.append(line.strip())

    return "\n".join(texts)

def iter_docs(directory: str) -> List[Tuple[str, str]]:
    out = []
    if not os.path.isdir(directory):
        return out

    for root, _, files in os.walk(directory):
        for fn in files:
            fp = os.path.join(root, fn)
            ext = os.path.splitext(fn)[1].lower()
            print(f"🔎 Found file: {fp} (ext={ext})")
            try:
                if ext == ".pdf":
                    text = load_pdf(fp)
                elif ext == ".docx":
                    text = load_docx(fp)
                elif ext == ".csv":
                    text = load_csv(fp)
                elif ext in (".txt", ".md"):
                    text = load_txt(fp)
                else:
                    print(f"⏩ Skipping unsupported file: {fp}")
                    continue

                if text.strip():
                    print(f"✅ Loaded {len(text)} chars from {fp}")
                    out.append((fp, text))
                else:
                    print(f"⚠️ Empty text extracted from {fp}")
            except Exception as e:
                print(f"❌ Failed to load {fp}: {e}")
    return out

# --------- Index management ---------
def create_empty_index(dim: int):
    return faiss.IndexFlatIP(dim)

def save_index(index, meta_rows: List[Dict]):
    os.makedirs(STORAGE_DIR, exist_ok=True)
    faiss.write_index(index, INDEX_PATH)
    with open(META_PATH, "w", encoding="utf-8") as f:
        for row in meta_rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")

def load_meta() -> List[Dict]:
    if not os.path.exists(META_PATH):
        return []
    with open(META_PATH, "r", encoding="utf-8") as f:
        return [json.loads(line) for line in f if line.strip()]

def rebuild_index() -> Dict:
    os.makedirs(STORAGE_DIR, exist_ok=True)
    docs = iter_docs(KNOWLEDGE_DIR)
    if not docs:
        return {"status": "no_files", "message": f"No supported files found in {KNOWLEDGE_DIR}."}

    all_chunks, meta_rows = [], []
    added_sources = set()

    for fp, text in docs:
        rel = os.path.relpath(fp, KNOWLEDGE_DIR)
        added_sources.add(rel)
        for i, ch in enumerate(chunk_text(text)):
            meta_rows.append({
                "id": str(uuid.uuid4()),
                "source": rel,
                "chunk_index": i,
                "text": ch
            })
            all_chunks.append(ch)

    if not all_chunks:
        return {"status": "no_chunks", "message": "No text chunks created."}

    embs = embed_texts(all_chunks)
    index = create_empty_index(dim=embs.shape[1])
    index.add(embs)
    save_index(index, meta_rows)

    resp = {
        "status": "ok",
        "chunks": len(all_chunks),
        "docs": len(added_sources),
        "sources": sorted(list(added_sources))
    }
    print(f"[reindex] {resp}")
    return resp

def search(query: str, k: int = TOP_K) -> List[Dict]:
    meta = load_meta()
    if (not os.path.exists(INDEX_PATH)) or (not meta):
        return []
    index = faiss.read_index(INDEX_PATH)
    q_emb = embed_texts([query])
    D, I = index.search(q_emb, min(k, index.ntotal))
    return [
        {
            "score": float(score),
            "source": meta[idx]["source"],
            "chunk_index": meta[idx]["chunk_index"],
            "text": meta[idx]["text"]
        }
        for score, idx in zip(D[0], I[0]) if idx != -1
    ]

# --------- LLM call (Groq) ---------
def call_groq(system_prompt: str, user_prompt: str) -> str:
    if not GROQ_API_KEY:
        return "[Missing GROQ_API_KEY] Please set your API key."
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.2,
        "max_tokens": 800
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=60)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]

# --------- Prompting ---------
BASE_SYSTEM_PROMPT = (
    "You are a helpful CSR process assistant for our company. "
    "Answer succinctly and ONLY based on the provided context. "
    "If the answer is not in the context, say you don't know and suggest who to ask."
)

def build_qa_prompt(question: str, context_chunks: List[Dict]) -> str:
    context = "\n\n".join([
        f"[Source: {c['source']} | Score: {c['score']:.3f}]\n{c['text']}" for c in context_chunks
    ])
    return f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"

# ==================== ROUTES ====================
@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")

@app.route("/api/reindex", methods=["POST"])
def api_reindex():
    return jsonify(rebuild_index())

@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.get_json(force=True)
    question = data.get("message", "").strip()
    if not question:
        return jsonify({"error": "Empty message"}), 400
    hits = search(question, k=TOP_K)
    if not hits:
        return jsonify({"reply": "No index found. Add docs and reindex."})
    try:
        answer = call_groq(BASE_SYSTEM_PROMPT, build_qa_prompt(question, hits))
    except Exception as e:
        answer = f"LLM call failed: {e}"
    return jsonify({"reply": answer, "sources": hits})

@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No selected file"}), 400
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in SUPPORTED_EXTS:
        return jsonify({"error": f"Unsupported file type: {ext}"}), 400

    os.makedirs(KNOWLEDGE_DIR, exist_ok=True)
    save_path = os.path.join(KNOWLEDGE_DIR, file.filename)
    file.save(save_path)
    return jsonify({"status": "uploaded", "path": save_path})

if __name__ == "__main__":
    print(f"[startup] KNOWLEDGE_DIR = {KNOWLEDGE_DIR}")
    print(f"[startup] STORAGE_DIR    = {STORAGE_DIR}")
    port = int(os.getenv("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=True, use_reloader=False)
